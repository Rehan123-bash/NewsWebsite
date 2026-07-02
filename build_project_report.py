from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT_PATH = ROOT / "News_Web_Project_Status_Report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_inches):
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))


def style_document(document: Document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(0x2E, 0x74, 0xB5)),
        ("Heading 2", 13, RGBColor(0x2E, 0x74, 0xB5)),
        ("Heading 3", 12, RGBColor(0x1F, 0x4D, 0x78)),
    ]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_title(document: Document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("News Web Project Status Report")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    meta = document.add_paragraph()
    meta.paragraph_format.space_after = Pt(10)
    meta_run = meta.add_run("Prepared from the current codebase scan on 30 June 2026")
    meta_run.font.size = Pt(10.5)
    meta_run.font.color.rgb = RGBColor(0x55, 0x65, 0x76)


def add_bullets(document: Document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(document: Document, items):
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_phase_table(document: Document):
    document.add_heading("Recommended Delivery Phases", level=1)
    intro = document.add_paragraph(
        "The project is best treated as a phased newsroom platform build. "
        "The current repository is already in the admin and CMS foundation stage."
    )
    intro.paragraph_format.space_after = Pt(8)

    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    headers = ["Phase", "Focus", "What happens here", "Status"]
    widths = [0.8, 1.45, 3.55, 0.7]
    for idx, (cell, text, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
        set_cell_width(cell, width)
        cell.text = text
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, "F2F4F7")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.runs[0]
        run.font.bold = True
        run.font.size = Pt(10.5)

    rows = [
        (
            "1",
            "CMS foundation",
            "Backend models, authentication, admin shell, article/category/tag/media structure, API contracts, and basic article creation.",
            "Mostly done",
        ),
        (
            "2",
            "Admin completion",
            "Connect all admin pages to real APIs, add edit/delete flows, build media management UI, wire search, and replace dashboard mock stats with live data.",
            "In progress",
        ),
        (
            "3",
            "User-facing news site",
            "Build homepage, category pages, article detail pages, trending/latest sections, header/footer, article listing cards, and public navigation.",
            "Not started",
        ),
        (
            "4",
            "Growth and polish",
            "SEO improvements, analytics, performance tuning, scheduling automation, richer editor experience, testing, and deployment hardening.",
            "Planned",
        ),
    ]

    for row in rows:
        cells = table.add_row().cells
        for idx, (cell, text, width) in enumerate(zip(cells, row, widths)):
            set_cell_width(cell, width)
            cell.text = text
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(10.5)


def build():
    document = Document()
    style_document(document)
    add_title(document)

    document.add_heading("1. Current Project Snapshot", level=1)
    document.add_paragraph(
        "This repository is currently structured as a two-part newsroom application: "
        "a Django REST backend for content management and a Next.js frontend that is focused on the internal admin side."
    )
    add_bullets(
        document,
        [
            "Backend apps already exist for accounts, articles, categories, tags, and media assets.",
            "JWT-based login, logout, refresh, registration, and current-user endpoints are present.",
            "Role checks are implemented for admin and editor access.",
            "The frontend already has a login screen, admin dashboard shell, and separate admin routes for articles, categories, tags, authors, and search.",
            "The root frontend route redirects to the login page, which means the public reader-facing site has not been started yet.",
        ],
    )

    document.add_heading("2. What Has Been Done", level=1)
    document.add_heading("Backend foundation", level=2)
    add_bullets(
        document,
        [
            "Custom user model with newsroom roles: admin, editor, and journalist.",
            "Article model includes slugging, status handling, publication timestamps, scheduling fields, featured/breaking flags, view count, reading time, and SEO metadata.",
            "Category and tag models auto-generate unique slugs.",
            "Media library supports file upload, captions, alt text, and uploader tracking.",
            "OpenAPI schema and Swagger docs endpoints are configured.",
            "CORS and database configuration are set up for local development and future PostgreSQL use.",
        ],
    )

    document.add_heading("Frontend foundation", level=2)
    add_bullets(
        document,
        [
            "Next.js app router project is scaffolded with Tailwind CSS and React Query.",
            "Axios client is configured to use the backend API with bearer tokens from local storage.",
            "A login page is already wired to the backend auth endpoint.",
            "Admin layout includes route protection, a reusable shell, navigation, and logout behavior.",
            "Article listing and new-article creation pages already call the real backend endpoints.",
        ],
    )

    document.add_heading("3. What Still Needs To Be Done", level=1)
    add_numbered(
        document,
        [
            "Finish the admin side by converting placeholder pages and mock dashboard content into real API-backed screens.",
            "Add missing CRUD flows such as edit article, delete article, create and manage tags/categories from the UI, and author management endpoints/screens.",
            "Build the media management frontend page because the sidebar links to media but no media page is present yet.",
            "Create public reader-facing pages including homepage, article detail page, category pages, and search results page for visitors.",
            "Add stronger publishing workflow features such as scheduled publishing UI, archive flows, live statistics, and moderation/editorial controls.",
            "Introduce testing, deployment setup, production database/storage configuration, and performance/SEO hardening.",
        ],
    )

    document.add_heading("4. Gaps and Observations From the Current Code", level=1)
    add_bullets(
        document,
        [
            "The admin dashboard currently shows hard-coded sample statistics and activity instead of live newsroom data.",
            "The tags page is still static sample content instead of using the tag API.",
            "The authors page is only a placeholder note and depends on profile endpoints that are not built yet.",
            "The search page UI exists, but the frontend search flow is not connected despite backend article filtering support already existing.",
            "Category and tag list endpoints currently require authentication, so public discovery pages will need either new public endpoints or adjusted permissions.",
            "Analytics and SEO Django apps are present in settings but do not yet contain working feature code.",
        ],
    )

    add_phase_table(document)

    document.add_heading("When To Build the User Side", level=1)
    document.add_paragraph(
        "The user-facing news website should begin in Phase 3, after the admin/CMS side is stable enough for editors to create, tag, "
        "categorize, and publish content reliably. Building the public site too early would make design work move faster than content operations, "
        "which usually creates rework."
    )
    add_bullets(
        document,
        [
            "Phase 1 and Phase 2 should make sure content can be created and managed properly.",
            "Phase 3 should focus on the reader experience: homepage, article pages, category landing pages, latest and trending sections, and responsive navigation.",
            "Once the public pages exist, Phase 4 can improve SEO, analytics, performance, and editorial polish.",
        ],
    )

    document.add_heading("Suggested Immediate Next Steps", level=1)
    add_numbered(
        document,
        [
            "Connect the admin dashboard cards and activity feed to live article, category, and user data.",
            "Build the missing media management page and complete tag/category management actions.",
            "Add edit/update flows for articles so the CMS is usable beyond just creation and listing.",
            "Decide the exact public news site structure and design system before starting Phase 3.",
            "After the admin workflow is stable, start the user-facing homepage and article detail pages first.",
        ],
    )

    document.add_heading("Conclusion", level=1)
    document.add_paragraph(
        "The project is not at the idea stage anymore. It already has a meaningful CMS base with authentication, content models, publishing logic, "
        "and an initial admin UI. The biggest remaining shift is moving from internal tooling to a complete product by finishing the admin workflows first "
        "and then building the public reader side in the next major phase."
    )

    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build()
